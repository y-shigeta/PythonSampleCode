import docx
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX
from os import path
import logging

OrgWord = "SampleWord.docx"
NewWord = "./Updated_"+OrgWord

#Initialize logging
LogfileName = "./"+path.splitext(__file__)[0]+".log"
formatter = '%(asctime)s : %(levelname)s : %(message)s'
logging.basicConfig(level=logging.DEBUG, format=formatter)
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

#initialize parameters
strCurrentHeader = []
cnt = 0
headerflag = 0


# Open Word file
doc = Document(OrgWord)

for parag in doc.paragraphs:
        cnt=cnt+1
        logger.debug('paragraph %s', parag.text)

        ### List Formatted Strings
        #if parag.text.find(SearchString) > -1:
        #        logger.info("[FOUND String!!]:%s :%s",str(cnt),parag.text)
        
        ### Find Indent right after Header2&3
        if headerflag == 2:
                logger.info("[FOUND Indent After Style Header2!!]:%s indent:%s :%s",strCurrentHeader,str(parag.paragraph_format.left_indent),parag.text)
        if headerflag == 3:
                logger.info("[FOUND Indent After Style Header3!!]:%s indent:%s :%s",strCurrentHeader,str(parag.paragraph_format.left_indent),parag.text)
        headerflag = 0

        ### Find Style
        if parag.style.name == "Heading 1": 
                headerflag = 1
                strCurrentHeader = parag.text
                logger.info("[FOUND Style Header1!!]:%s fontsize:%s :%s",strCurrentHeader,str(parag.style.font.size),parag.text)
        elif parag.style.name == "Heading 2": 
                headerflag = 2
                strCurrentHeader = parag.text
                logger.info("[FOUND Style Header2!!]:%s fontsize:%s :%s",strCurrentHeader,str(parag.style.font.size),parag.text)
        elif parag.style.name == "Heading 3": 
                headerflag = 3
                strCurrentHeader = parag.text
                logger.info("[FOUND Style Header3!!]:%s fontsize:%s :%s",strCurrentHeader,str(parag.style.font.size),parag.text)
        else:
                for run in parag.runs:
#                       logger.debug(run.font.name),str(run.fontsize
                        strText = run.text
                        strFontName = run.font.name
                        strFontSize = run.font.size
                        
                        logger.info("FontName:%s FontSize:%s Text:%s",strFontName,strFontSize,strText)
                        
#                        logger.info("[FOUND Fontsize is NOT 10.5!!]:%s fontsize:%s :%s",strCurrentHeader,strFontName,strFontSize,parag.text)
#                       run.font.name

#                if parag.text.fontsize != "10.5" or parag.text.font  :
#                        logger.info("[FOUND Fontsize is NOT 10.5!!]:%s fontsize:%s :%s",strCurrentHeader,str(parag.style.font.size),parag.text)


#parag = doc.add_paragraph("Left indent test")
#parag.left_indent = Inches(.25)
#doc.parag[0].runs[0].font.color.rpb = WD_COLOR_INDEX.YELLOW

doc.save(NewWord)
