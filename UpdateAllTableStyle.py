import docx
from docx import Document
from docx.shared import Inches
from os import path
import logging

OrgWord = "SampleWord.docx"
NewWord = "./Updated_"+OrgWord

#Initialize logging
LogfileName = "./"+path.splitext(__file__)[0]+".log"
formatter = '%(asctime)s : %(levelname)s : %(message)s'
logging.basicConfig(level=logging.DEBUG, format=formatter,filename=LogfileName)
#logging.basicConfig(level=logging.DEBUG, format=formatter)
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# Open Word file
doc = Document(OrgWord)

### Logging all Table Stylesheet before update
logger.info("### All Table StyleSheet Information ###")
for tbl in doc.tables:
        logger.info("Table StyleSheet: %s ",tbl.style.name)

### Update all table style sheets with the first table style sheet
i = 0
while i != len(doc.tables):
        if i == 0:
                FirstTableStyle = doc.tables[0].style
        else:
                doc.tables[i].style = FirstTableStyle
        i += 1
        continue

doc.save(NewWord)
